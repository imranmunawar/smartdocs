from docx import Document
from docx.shared import Inches
from django.conf import settings
import os
from docx import Document
from python_docx_replace import docx_replace,docx_get_keys
# from ..s3.client import S3_Client
from ..gcs.client import GCS_Client

class ProcessDocument():

    def process_and_save(self, inputMap:dict, imageMap:dict, inputPath:str, outputPath:str):

        # s3 = S3_Client()
        # doc_content = s3.read_document_from_s3(inputPath)
        gcs = GCS_Client()
        doc_content = gcs.read_document_from_gcs(inputPath)

        if doc_content is None:
            print("Failed to read document from S3.")
            return False

        doc = Document(doc_content)
        for key, value in inputMap.items():
            self.edit_docx(doc, key, value, font_type='Times New Roman')
        for key, value in imageMap.items():
            self.replace_placeholder_with_image(doc, key, value)

        doc.save(outputPath)
        self.second_script(inputMap, outputPath)

        return True


    def second_script(self, inputMap, outputPath):
        newMap = self.update_placeholders(inputMap)
        doc = Document(outputPath)
        docx_replace(doc, **newMap)
        doc.save(outputPath)


    def update_placeholders(self, inputMap):
        newMap = {}

        for key, value in inputMap.items():
            new_key = key.replace("${", "").replace("}", "")
            newMap[new_key] = value

        return newMap


    def replace_text_in_run(self, run, old_text, new_text, font_type):
        if old_text in run.text:
            original_font_size = run.font.size
            run.font.name = font_type
            run.text = run.text.replace(old_text, new_text)
            run.font.size = original_font_size

    def combine_text_run(self, all_runs):
        text_runs = ''
        for run in all_runs:
            text_runs += run.text

        return text_runs

    def edit_docx(self, doc, old_text, new_text, font_type='Times New Roman'):
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                combine_text = self.combine_text_run(paragraph.runs)

                if old_text == combine_text:
                    index = 0
                    for run in paragraph.runs:
                        original_font_size = run.font.size
                        run.font.name = font_type
                        if index == (len(paragraph.runs) - 1):
                            run.text = new_text
                        else:
                            run.text = ""
                        run.font.size = original_font_size
                        index = index + 1
                else:
                    for run in paragraph.runs:
                        self.replace_text_in_run(run, old_text, new_text, font_type)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        for run in cell.paragraphs[0].runs:
                            self.replace_text_in_run(run, old_text, new_text, font_type)

        txbx = doc.inline_shapes._body.xpath('//w:txbxContent')
        for i in range(0, len(txbx)):
            for tx_idx, tx in enumerate(txbx[i]):
                children = tx.getchildren()
                for child_idx, child in enumerate(children):
                    if child.text and old_text in child.text:
                        child.text= child.text.replace(old_text, new_text)
    
    def replace_placeholder_with_image(self, doc, placeholder_text, file_name):

        image_path = os.path.join(settings.IMAGE_FILES, file_name)
        # s3 = S3_Client()
        # image_content = s3.download_image_from_s3(file_name)

        gcs = GCS_Client()
        image_content = gcs.download_image_from_gcs(file_name)
        
        if image_content is None:
            print(f"Could not download image {file_name} from S3")
            return

        with open(image_path, 'wb') as image_file:
            image_file.write(image_content)

        try:
            for paragraph in doc.paragraphs:
                if placeholder_text in paragraph.text:
                    run = paragraph.runs[0]
                    run.clear()
                    run.add_picture(image_path)
        finally:
            os.remove(image_path)
